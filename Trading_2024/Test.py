#!/usr/bin/env python3
"""
strip_images.py
Remove embedded images from:
  - MS Word .docx (python-docx)
  - PDF .pdf (PyMuPDF / fitz)

Behavior:
- DOCX: removes drawing/pict nodes + deletes image relationships.
- PDF : redacts rectangles where raster images appear (white fill).
  WARNING: scanned PDFs (page is an image) will turn blank.

Usage:
  python strip_images.py input.docx
  python strip_images.py input.pdf
  python strip_images.py /path/to/folder --recursive
  python strip_images.py /path/to/folder --recursive --outdir /path/out
"""

import argparse
import sys
from pathlib import Path
from typing import Iterable, Optional

from docx import Document

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None


# ---------------- DOCX ----------------

def iter_all_paragraphs(doc: Document) -> Iterable:
    """Yield all paragraphs including those inside tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def strip_docx_images(in_path: Path, out_path: Path) -> None:
    doc = Document(str(in_path))

    # 1) Remove drawing/pict elements from runs without string-serializing XML
    # Word images are typically under w:r -> w:drawing (DrawingML) or w:pict (VML)
    removed_nodes = 0
    removed_runs = 0

    for p in iter_all_paragraphs(doc):
        for run in list(p.runs):
            r = run._element  # CT_R
            # find drawing/pict children
            to_remove = []
            for child in list(r):
                tag = child.tag
                if tag.endswith("}drawing") or tag.endswith("}pict"):
                    to_remove.append(child)

            for child in to_remove:
                r.remove(child)
                removed_nodes += 1

            # Optional: if run becomes empty (no text and no children), remove run
            # (prevents lingering empty runs that can affect spacing)
            if (len(r) == 0) and (run.text == ""):
                p._p.remove(r)
                removed_runs += 1

    # 2) Remove image relationships so word/media binaries are dropped
    rels = doc.part._rels
    removed_rels = 0
    for rel_id in list(rels.keys()):
        rel = rels[rel_id]
        target_ref = (getattr(rel, "target_ref", "") or "").lower()
        if "media/" in target_ref and "image" in target_ref:
            del rels[rel_id]
            removed_rels += 1

    doc.save(str(out_path))
    print(f"[DOCX] removed_nodes={removed_nodes} removed_runs={removed_runs} removed_rels={removed_rels}")


# ---------------- PDF ----------------

def strip_pdf_images(in_path: Path, out_path: Path, fill_white: bool = True) -> None:
    if fitz is None:
        raise RuntimeError("PyMuPDF not installed. Install: pip install pymupdf")

    doc = fitz.open(str(in_path))

    total_rects = 0
    total_pages = 0

    for page in doc:
        total_pages += 1
        imgs = page.get_images(full=True)
        if not imgs:
            continue

        rects = []
        for img in imgs:
            xref = img[0]
            rects.extend(page.get_image_rects(xref))

        if not rects:
            continue

        # Deduplicate rectangles (PyMuPDF rect is hashable via tuple)
        uniq = []
        seen = set()
        for r in rects:
            key = (round(r.x0, 3), round(r.y0, 3), round(r.x1, 3), round(r.y1, 3))
            if key not in seen:
                seen.add(key)
                uniq.append(r)

        for r in uniq:
            if fill_white:
                page.add_redact_annot(r, fill=(1, 1, 1))
            else:
                page.add_redact_annot(r)

        page.apply_redactions()
        page.clean_contents()  # helps shrink after removals
        total_rects += len(uniq)

    doc.save(str(out_path), garbage=4, deflate=True)
    doc.close()
    print(f"[PDF] pages={total_pages} redacted_rects={total_rects}")


# ---------------- IO / CLI ----------------

def unique_out_path(out_path: Path) -> Path:
    """If output exists, add _v2/_v3..."""
    if not out_path.exists():
        return out_path
    stem, suffix = out_path.stem, out_path.suffix
    i = 2
    while True:
        cand = out_path.with_name(f"{stem}_v{i}{suffix}")
        if not cand.exists():
            return cand
        i += 1


def default_out_path(in_path: Path, outdir: Optional[Path]) -> Path:
    suffix = in_path.suffix.lower()
    stem = in_path.stem
    if stem.endswith("_noimg"):
        out_name = f"{stem}{suffix}"  # don't stack _noimg_noimg
    else:
        out_name = f"{stem}_noimg{suffix}"
    base = (outdir / out_name) if outdir else (in_path.parent / out_name)
    return unique_out_path(base)


def process_file(fp: Path, outdir: Optional[Path]) -> bool:
    suffix = fp.suffix.lower()
    out_path = default_out_path(fp, outdir)

    if suffix == ".docx":
        strip_docx_images(fp, out_path)
        print(f"[OK] DOCX -> {out_path}")
        return True
    if suffix == ".pdf":
        strip_pdf_images(fp, out_path)
        print(f"[OK] PDF  -> {out_path}")
        return True
    return False


def main():
    ap = argparse.ArgumentParser(description="Remove images from DOCX and PDF files.")
    ap.add_argument("--input", default=r"C:\Users\himan\Documents\Audacity", help="Input file or folder path")
    ap.add_argument("--recursive", action="store_true", help="Recurse into subfolders")
    ap.add_argument("--outdir", default=r"C:\Users\himan\Documents", help="Output directory (optional)")
    args = ap.parse_args()

    in_path = Path(args.input).expanduser().resolve()
    outdir = Path(args.outdir).expanduser().resolve() if args.outdir else None
    if outdir:
        outdir.mkdir(parents=True, exist_ok=True)

    processed = 0
    failed = 0

    def handle(fp: Path):
        nonlocal processed, failed
        try:
            if process_file(fp, outdir):
                processed += 1
        except Exception as e:
            failed += 1
            print(f"[FAIL] {fp}: {e}", file=sys.stderr)

    if in_path.is_file():
        handle(in_path)
    elif in_path.is_dir():
        pattern = "**/*" if args.recursive else "*"
        for fp in in_path.glob(pattern):
            if fp.is_file():
                handle(fp)
    else:
        print(f"[ERROR] Not found: {in_path}", file=sys.stderr)
        sys.exit(2)

    print(f"[DONE] processed={processed} failed={failed}")
    if failed:
        sys.exit(1)


if __name__ == "__main__":
    main()